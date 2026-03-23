#!/usr/bin/env python3
"""
InkDrop — DOCX Signature Insertion Engine
Drops a processed signature PNG into a Word document.

Supports:
  - Keyword-based placement (finds "Signature:", "Sign here", "___")
  - End-of-document placement (appends after last paragraph)
  - Paragraph index placement (inserts after specific paragraph)

Usage:
    python insert_docx.py document.docx signature.png output.docx --find "Signature:"
    python insert_docx.py document.docx signature.png output.docx --after-paragraph 12
    python insert_docx.py document.docx signature.png output.docx --end
"""

import argparse
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is required.")
    print("  pip install python-docx --break-system-packages")
    sys.exit(1)


def find_signature_paragraph(doc: Document, keyword: str) -> int | None:
    """Find the paragraph index containing the keyword."""
    patterns = [re.escape(keyword)]
    patterns.extend([r"_{5,}", r"\.{5,}", r"-{5,}"])

    for i, para in enumerate(doc.paragraphs):
        text = para.text
        for pattern in patterns:
            if re.search(pattern, text, re.IGNORECASE):
                return i
    return None


def insert_signature_docx(
    docx_path: str,
    sig_path: str,
    output_path: str,
    keyword: str = None,
    after_paragraph: int = None,
    end: bool = False,
    sig_width_inches: float = 2.0,
    add_line: bool = False,
    add_name: str = None,
    add_date: bool = False,
) -> dict:
    """
    Insert a signature image into a DOCX document.

    Placement modes (in priority order):
        1. keyword          — search for text marker
        2. after_paragraph  — insert after paragraph index
        3. end              — append to end of document

    Args:
        docx_path:        Source DOCX
        sig_path:         Processed signature PNG
        output_path:      Output DOCX path
        keyword:          Text to search for placement
        after_paragraph:  Paragraph index to insert after
        end:              Append to end of document
        sig_width_inches: Signature width in inches (default 2.0)
        add_line:         Add a line below the signature
        add_name:         Add printed name below signature
        add_date:         Add date below signature

    Returns:
        dict with placement metadata
    """
    doc = Document(docx_path)
    target_idx = None

    if keyword:
        target_idx = find_signature_paragraph(doc, keyword)
        if target_idx is None:
            raise ValueError(f"Keyword '{keyword}' not found in document")
    elif after_paragraph is not None:
        target_idx = after_paragraph
    elif end:
        target_idx = len(doc.paragraphs) - 1
    else:
        # Default: end of document
        target_idx = len(doc.paragraphs) - 1

    # Insert signature image
    # We add a new paragraph after the target and insert the image there
    target_para = doc.paragraphs[target_idx]

    # Add signature image paragraph
    sig_para = doc.add_paragraph()
    sig_run = sig_para.add_run()
    sig_run.add_picture(sig_path, width=Inches(sig_width_inches))

    # Move the new paragraph to right after the target
    target_para._element.addnext(sig_para._element)

    # Optional: add line below signature
    if add_line:
        line_para = doc.add_paragraph("_" * 40)
        line_para.style.font.size = Pt(10)
        sig_para._element.addnext(line_para._element)

        # Optional: printed name
        if add_name:
            name_para = doc.add_paragraph(add_name)
            name_para.style.font.size = Pt(10)
            line_para._element.addnext(name_para._element)

    # Optional: add date
    if add_date:
        from datetime import datetime
        date_str = datetime.now().strftime("%B %d, %Y")
        date_para = doc.add_paragraph(f"Date: {date_str}")
        date_para.style.font.size = Pt(10)
        # Insert after the last element we added
        last = sig_para._element
        for sibling in sig_para._element.itersiblings():
            last = sibling
        last.addnext(date_para._element)

    doc.save(output_path)

    return {
        "paragraph_index": target_idx,
        "matched_text": target_para.text[:60] if target_para.text else "(empty)",
        "sig_width_inches": sig_width_inches,
        "output": output_path,
    }


def main():
    parser = argparse.ArgumentParser(description="InkDrop — Insert signature into DOCX")
    parser.add_argument("docx", help="Source DOCX file")
    parser.add_argument("signature", help="Processed signature PNG")
    parser.add_argument("output", help="Output DOCX path")
    parser.add_argument("--find", type=str, help="Keyword to locate signature line")
    parser.add_argument("--after-paragraph", type=int, help="Insert after paragraph index")
    parser.add_argument("--end", action="store_true", help="Append to end of document")
    parser.add_argument("--width", type=float, default=2.0, help="Signature width in inches")
    parser.add_argument("--line", action="store_true", help="Add signature line below")
    parser.add_argument("--name", type=str, help="Add printed name below signature")
    parser.add_argument("--date", action="store_true", help="Add current date")

    args = parser.parse_args()

    for f in [args.docx, args.signature]:
        if not Path(f).exists():
            print(f"ERROR: File not found: {f}")
            sys.exit(1)

    result = insert_signature_docx(
        args.docx, args.signature, args.output,
        keyword=args.find, after_paragraph=args.after_paragraph,
        end=args.end, sig_width_inches=args.width,
        add_line=args.line, add_name=args.name, add_date=args.date,
    )

    print(f"✓ InkDrop signed: paragraph {result['paragraph_index']}")
    print(f"  Matched: \"{result['matched_text']}\"")
    print(f"  Output: {result['output']}")


if __name__ == "__main__":
    main()
